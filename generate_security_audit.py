"""Generate Wisper Secure Build Checklist + Security Audit Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
DOCX_OUT = ROOT / "Wisper-Security-Audit.docx"

# id, severity, category, location, description, recommendation
FINDINGS = [
    (
        "SEC-001",
        "High",
        "Arbitrary file write (CWE-73)",
        "wisper/src-tauri/src/lib.rs — write_text_file",
        "The write_text_file IPC command writes to any path the frontend supplies. "
        "There is no canonicalization, extension check, or directory allowlist. "
        "A modified frontend or devtools invoke could overwrite sensitive files.",
        "Restrict writes to user-selected save-dialog paths only; canonicalize the path; "
        "reject paths outside an allowlist (e.g. Downloads/Documents).",
    ),
    (
        "SEC-002",
        "High",
        "SSRF / internal network probe (CWE-918)",
        "wisper/wisper-core/src/fetch/mod.rs — normalize_url, download_url",
        "normalize_url only checks for http:// or https:// prefix. URLs such as "
        "http://127.0.0.1:8080/admin or http://169.254.169.254/ pass validation and are "
        "handed to yt-dlp, which fetches them from the user's machine.",
        "Block private, loopback, and link-local IP ranges; parse URLs with a proper crate; "
        "consider hostname allowlists for known media hosts.",
    ),
    (
        "SEC-003",
        "Medium",
        "Unrestricted file read via IPC (CWE-22)",
        "wisper/src-tauri/src/lib.rs — start_transcription",
        "start_transcription accepts any audio_path string with no check that the path is "
        "user-selected, under app data, or a regular audio file.",
        "Validate path exists, is a file, has an allowed extension, and optionally lies "
        "under the app audio directory or a user-picked path registry.",
    ),
    (
        "SEC-004",
        "Medium",
        "Missing Content Security Policy (CWE-1021)",
        "wisper/src-tauri/tauri.conf.json",
        "csp is set to null, disabling CSP in the production webview. "
        "If XSS were introduced in the React UI, there is no CSP backstop.",
        "Define a strict production CSP (default-src 'self'; minimize unsafe-inline).",
    ),
    (
        "SEC-005",
        "Medium",
        "Over-broad Tauri capabilities (CWE-250)",
        "wisper/src-tauri/capabilities/default.json; lib.rs",
        "Grants core:default, dialog:default, and opener:default. "
        "The opener plugin is registered but not used in App.tsx — unused attack surface.",
        "Remove unused opener plugin; replace core:default with explicit minimal permissions.",
    ),
    (
        "SEC-006",
        "Medium",
        "Dev dependency CVEs (CWE-1395)",
        "wisper/package-lock.json (vite / esbuild chain)",
        "npm audit reports 3 High severity issues in the dev toolchain (esbuild via vite). "
        "These affect npm run dev, not the shipped Tauri bundle.",
        "Track vite/esbuild upgrades; run npm audit in CI; avoid exposing dev server to LAN.",
    ),
    (
        "SEC-007",
        "Medium",
        "PII in debug logs (CWE-532)",
        "wisper/wisper-core/src/transcribe/mod.rs; audio/mod.rs",
        "eprintln! logs full filesystem paths and decode metadata to stderr. "
        "Paths can reveal usernames and sensitive folder names. Transcript text is not logged.",
        "Gate diagnostic logging behind a debug flag; redact home-directory prefixes in release builds.",
    ),
    (
        "SEC-008",
        "Medium",
        "Model download integrity (CWE-494)",
        "wisper/scripts/download-model.ps1",
        "Whisper models are downloaded from Hugging Face with no checksum or signature verification.",
        "Pin SHA-256 hashes; verify before use; document expected file size.",
    ),
    (
        "SEC-009",
        "Low",
        "App data fallback to CWD (CWE-22)",
        "wisper/src-tauri/src/lib.rs — app_data_dir",
        "If Tauri path resolution fails, app_data_dir falls back to '.' — placing DB, models, "
        "and audio in the process working directory.",
        "Fail fast if app_data_dir is unavailable; never fall back to the current directory.",
    ),
    (
        "SEC-010",
        "Low",
        "Server-side language input not validated (CWE-20)",
        "wisper/src-tauri/src/lib.rs",
        "The language parameter from IPC is passed to Whisper after trim/lowercase only. "
        "The frontend whitelists options, but IPC can be invoked directly.",
        "Validate language against the same whitelist as LANGUAGE_OPTIONS in Rust.",
    ),
    (
        "SEC-011",
        "Low",
        "Resource exhaustion / DoS (CWE-400)",
        "wisper/wisper-core/src/audio/mod.rs; transcribe/mod.rs",
        "No file-size or duration cap on import, decode, or transcribe. "
        "Very large or malicious media can exhaust RAM or CPU.",
        "Enforce max file size and max audio duration before decode.",
    ),
    (
        "SEC-012",
        "Low",
        "CI action pinning inconsistency (CWE-829)",
        ".github/workflows/desktop.yml; release.yml",
        "Some actions are pinned to SHAs (setup-node, cuda-toolkit); others use floating tags "
        "(actions/checkout@v4, upload-artifact@v4, softprops/action-gh-release@v2).",
        "Pin all third-party GitHub Actions to full commit SHAs.",
    ),
    (
        "SEC-013",
        "Info",
        "No automated secret or dependency scanning",
        ".github/workflows/",
        "No Dependabot, CodeQL, gitleaks, or cargo audit in CI.",
        "Add Dependabot; run cargo audit and npm audit in the desktop workflow.",
    ),
    (
        "SEC-014",
        "Info",
        "Safe subprocess pattern (positive)",
        "fetch/mod.rs; audio/mod.rs",
        "yt-dlp and ffmpeg use Command::new().args([...]) without a shell — no shell injection "
        "from path or URL interpolation.",
        "Keep this pattern; never use sh -c or cmd /C with user input.",
    ),
    (
        "SEC-015",
        "Info",
        "SQLite injection resistance (positive)",
        "wisper/wisper-core/src/storage/mod.rs",
        "All SQL uses params![] bound parameters. FTS search escapes double quotes before quoting.",
        "Keep parameterized queries; add FTS fuzz tests for edge-case queries.",
    ),
]

CHECKLIST = [
    "Secrets hygiene: confirm .env, API keys, and tokens are gitignored and never committed; run a secret scan before each release.",
    "Dependency audit: run npm audit (frontend) and cargo audit (Rust) in CI; pin or patch High/Critical findings.",
    "Action pinning: pin all GitHub Actions to immutable commit SHAs, not @v4 tags.",
    "Tauri CSP: set a production CSP in tauri.conf.json (default-src 'self'); verify the UI still works.",
    "Capability least privilege: remove unused opener:default; replace core:default with explicit permissions.",
    "IPC path validation: validate audio_path (allowed extensions, exists, is file) before decode/transcribe.",
    "IPC write validation: restrict write_text_file to user-confirmed save paths; reject writes outside intended directories.",
    "URL import hardening: block private, loopback, and link-local targets; document that URL import is the only network egress.",
    "Subprocess safety: keep Command::args() (no shell); pin/document yt-dlp and ffmpeg versions.",
    "Model integrity: verify Whisper model SHA-256 after download; reject unexpected files in models/.",
    "Logging policy: disable or redact eprintln! path diagnostics in release builds; transcripts must not appear in logs.",
    "Resource limits: cap import file size and max transcription duration to prevent DoS from crafted media.",
    "Release artifacts: sign Windows/macOS binaries; document that models and yt-dlp are user-installed dependencies.",
    "Privacy review: confirm UI accurately states when network is used (download only) vs offline (transcribe).",
    "Manual QA: test malicious inputs — odd URLs, huge files, path traversal strings, FTS special characters.",
]

POSITIVE = [
    "No cloud STT — transcription stays in-process via whisper.cpp; no outbound transcription API calls.",
    "No secrets in repository — grep found no API keys or tokens; .env is gitignored.",
    "Safe subprocess spawning — yt-dlp URL and ffmpeg paths passed as discrete args, not shell-joined.",
    "Parameterized SQLite — inserts, updates, and searches use bound parameters throughout storage.",
    "Scoped audio deletion — delete_recording only removes files under audio_root via path.starts_with.",
    "Partial download cleanup — PartialDownloadGuard removes yt-dlp partials on cancel or failure.",
    "URL scheme gate — non-HTTP(S) schemes are rejected in normalize_url.",
    "Frontend extension filter — drag/drop and file picker restrict to allowed audio/video extensions.",
    "React text rendering — no dangerouslySetInnerHTML; transcript rendered as text nodes.",
    "CI scoped permissions — release job requests only contents: write; no hardcoded secrets in workflows.",
]

OPEN_QUESTIONS = [
    "Are release MSI, DMG, and AppImage artifacts code-signed or notarized?",
    "Does Tauri 2 apply any default CSP when csp is null in a release build?",
    "Which yt-dlp versions were tested? Are there known CVEs in the user-installed version?",
    "Are bundled yt-dlp paths under resource_dir/bin/ ever populated in release builds?",
    "Run cargo audit on Rust dependencies — transitive CVE status not verified in this audit.",
    "macOS and Linux mic permission prompts and sandbox behavior need device testing.",
    "FTS edge cases: manual test search queries with quotes, wildcards, OR, and Unicode.",
    "Is the webview fully trusted? If not, every #[tauri::command] handler needs server-side validation.",
    "No runtime integrity check on whisper.cpp native library built via whisper-rs.",
    "Confirm no secrets in git history (gitleaks or git log -p) before public launch.",
]

INPUT_SURFACE = [
    ("Paste URL", "start_url_import", "yt-dlp subprocess (network + filesystem)"),
    ("Drag/drop or file picker path", "start_transcription", "symphonia/ffmpeg decode + Whisper"),
    ("Save dialog path", "write_text_file", "Arbitrary filesystem write"),
    ("Library search query", "search_library", "SQLite FTS5 MATCH"),
]


def set_cell_shading(cell, fill="D9E2F3"):
    """Light header shading — optional, skip if unsupported."""
    pass


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc():
    doc = Document()

    # Title page
    title = doc.add_heading("Wisper Security Audit Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Secure Build Checklist — L2 Clone Project")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Aisling · Wisper (local-first Whisper Notes clone)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Audit date: June 2026")
    doc.add_paragraph("Scope: Tauri 2 + React + Rust (wisper-core), yt-dlp, ffmpeg, SQLite, GitHub Actions")
    doc.add_paragraph("Status: Research and checklist draft — remediation deferred")
    doc.add_page_break()

    # Executive summary
    add_heading(doc, "1. Executive summary", 1)
    doc.add_paragraph(
        "Wisper is architecturally aligned with its privacy promise: transcription runs on-device via "
        "whisper.cpp, there are no cloud speech-to-text API keys in the repository, and SQLite uses "
        "parameterized queries. No Critical-severity findings were identified in source review."
    )
    doc.add_paragraph(
        "The main security gaps are trust-boundary issues, not secret leakage. Tauri IPC commands accept "
        "unvalidated file paths, URL import has minimal validation before yt-dlp runs, and Tauri hardening "
        "(Content Security Policy, capability scoping) is incomplete. The two highest practical risks are "
        "arbitrary file write via write_text_file and server-side request forgery (SSRF) via user-supplied URLs."
    )

    # Severity counts
    add_heading(doc, "2. Finding summary by severity", 1)
    counts = {}
    for f in FINDINGS:
        sev = f[1]
        if sev not in ("Info",):
            counts[sev] = counts.get(sev, 0) + 1
    st = doc.add_table(rows=2, cols=4)
    st.style = "Table Grid"
    for i, sev in enumerate(["High", "Medium", "Low", "Info"]):
        st.rows[0].cells[i].text = sev
        st.rows[1].cells[i].text = str(counts.get(sev, 0) if sev != "Info" else sum(1 for f in FINDINGS if f[1] == "Info"))

    doc.add_paragraph()

    # Course failure modes
    add_heading(doc, "3. L2 course — three failure modes", 1)
    doc.add_paragraph(
        "Every professional product team maintains a security checklist. This audit maps Wisper against "
        "three common failure modes when builders skip security thinking."
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Failure mode"
    h[1].text = "Wisper risk level"
    h[2].text = "Why"
    data = [
        (
            "1. API keys exposed in public repos",
            "Low",
            "No API keys found. .env is gitignored. No cloud STT — nothing to leak for transcription.",
        ),
        (
            "2. User data sent where it shouldn't go",
            "Medium",
            "Transcription stays local. yt-dlp makes outbound requests on URL import (by design). "
            "stderr logs can expose filesystem paths.",
        ),
        (
            "3. Products manipulated by unexpected inputs",
            "Highest",
            "IPC accepts paths and URLs without strong validation. Subprocess and filesystem sinks "
            "are reachable from user-controlled input.",
        ),
    ]
    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    add_heading(doc, "4. Most relevant failure mode for Wisper today", 1)
    doc.add_paragraph(
        "Failure mode #3 — products manipulated by unexpected inputs — is the most relevant to the clone today."
    )
    doc.add_paragraph(
        "Wisper's threat model is local processing, not API key exposure. The attack surface is "
        "user-controlled inputs crossing into privileged Rust code:"
    )
    it = doc.add_table(rows=1 + len(INPUT_SURFACE), cols=3)
    it.style = "Table Grid"
    it.rows[0].cells[0].text = "User input"
    it.rows[0].cells[1].text = "IPC command"
    it.rows[0].cells[2].text = "Privileged sink"
    for i, row in enumerate(INPUT_SURFACE, start=1):
        for j, val in enumerate(row):
            it.rows[i].cells[j].text = val

    # Detailed findings
    add_heading(doc, "5. Detailed findings", 1)
    doc.add_paragraph(
        "Each finding includes location, description, and recommended fix. Fixes are out of scope for "
        "this draft and will be addressed in a follow-up engineering pass."
    )
    for fid, sev, cat, loc, desc, rec in FINDINGS:
        add_heading(doc, f"{fid} — {sev}: {cat}", 2)
        doc.add_paragraph(f"Location: {loc}")
        doc.add_paragraph(f"Description: {desc}")
        doc.add_paragraph(f"Recommendation: {rec}")

    # Positive controls
    add_heading(doc, "6. Security controls already in place", 1)
    add_bullets(doc, POSITIVE)

    # Checklist
    add_heading(doc, "7. Secure Build Checklist (draft)", 1)
    doc.add_paragraph(
        "Use this checklist before each beta or public release. Add items across the L2 cycle as you "
        "learn what secure building looks like in practice."
    )
    for i, item in enumerate(CHECKLIST, start=1):
        doc.add_paragraph(f"☐ {i}. {item}")

    # Open questions
    add_heading(doc, "8. Open questions — manual verification needed", 1)
    add_bullets(doc, OPEN_QUESTIONS)

    # Teaching note
    add_heading(doc, "9. Teaching note for L2", 1)
    doc.add_paragraph(
        "Wisper is a strong example of privacy-by-architecture (local speech-to-text), but "
        "local-first does not mean input-safe. Subprocess boundaries (yt-dlp, ffmpeg) and Tauri IPC "
        "are where unexpected inputs become filesystem, network, or resource abuse — even without "
        "any API keys in the repository."
    )

    add_heading(doc, "10. Next steps", 1)
    doc.add_paragraph(
        "This document is audit and checklist only. Planned follow-up work includes IPC path validation, "
        "URL blocklists for private networks, production CSP, Tauri capability tightening, CI dependency "
        "scanning, and model checksum verification."
    )

    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    build_doc()
