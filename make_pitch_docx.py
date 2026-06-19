"""Generate Wisper pitch script as Word document."""
from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

OUT = Path(__file__).resolve().parent / "docs" / "Wisper_Pitch_2-3min.docx"
OUT_DOWNLOADS = Path(r"C:\Users\Aisling Ld Pursuit\Downloads\Wisper_Pitch_2-3min.docx")


def ensure_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def add_speaker(doc, name: str, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)


def main() -> None:
    ensure_docx()
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Wisper — 2–3 Minute Pitch", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Aisling Ld Pursuit & Jimmy")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Total time: ~2 min 30 sec  ·  Slides: 3  ·  Version: beta.21").font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading("Slide 1 — The Problem (~45 sec)", level=1)
    add_speaker(
        doc,
        "Aisling",
        "Imagine you're in a lecture or a client meeting. The speaker is moving fast, and you're trying to listen and take notes at the same time. You can't—people speak about three times faster than they type. So details get lost: names, numbers, action items.",
    )
    add_speaker(
        doc,
        "Jimmy",
        "And if you reach for a transcription app, most of them want your audio in the cloud. That's a non-starter for anything sensitive—therapy, legal, HR, research.",
    )
    add_speaker(
        doc,
        "Aisling",
        "There's Whisper Notes on Mac, which showed people want high-quality local transcription. But Windows and Linux users are still stuck choosing between bad notes and giving up their privacy.",
    )

    doc.add_heading("Slide 2 — The Solution (~60 sec)", level=1)
    add_speaker(
        doc,
        "Jimmy",
        "That's why we built Wisper—a cross-platform desktop app for Windows, Mac, and Linux.",
    )
    add_speaker(
        doc,
        "Aisling",
        "The flow is simple. First, capture audio—record from your mic, drop in a file, or paste a URL. On first launch, a welcome guide walks you through downloading a speech model and even recommends Small, Medium, or Large based on your hardware.",
    )
    add_speaker(
        doc,
        "Jimmy",
        "Second, transcribe locally. Everything runs on your machine using Whisper—we support CPU and GPU, including CUDA on NVIDIA, Metal on Mac, and Vulkan on AMD and Intel. Nothing leaves your computer. We say that right on the screen.",
    )
    add_speaker(
        doc,
        "Aisling",
        "Third, search and export. Every transcript goes into a local library with full-text search. You can edit segments and export to text. No account. No subscription. No upload.",
    )

    doc.add_heading("Slide 3 — How It's Built (~35 sec)", level=1)
    add_speaker(
        doc,
        "Jimmy",
        "Under the hood: React UI in Tauri 2, Rust core with whisper.cpp, SQLite for the library, and GPU backends where your machine supports them.",
    )
    add_speaker(
        doc,
        "Aisling",
        "We're at beta.21 today—all three model tiers ship, tier-aware transcription works, and we've been dogfooding on Windows with CUDA. Installers are on GitHub Releases for beta testers.",
    )

    doc.add_heading("Close (~10 sec)", level=1)
    add_speaker(
        doc,
        "Together",
        "Wisper is private, local, and fast enough for real work. We'd love your feedback—and if you have a meeting tomorrow, we'd rather you try it than take notes by hand.",
    )
    p = doc.add_paragraph()
    p.add_run("Optional if asked: ").bold = True
    p.add_run("GitHub — aislingld-pursuit/L2-Clone-Prodject · tag v0.2.0-beta.21")

    doc.add_heading("Presenter split", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Who leads"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    rows = [
        ("Problem (typing vs speech)", "Aisling"),
        ("Problem (cloud privacy)", "Jimmy"),
        ("Problem (market gap)", "Aisling"),
        ("Product intro + capture step", "Jimmy"),
        ("Transcribe + privacy", "Aisling"),
        ("Library + export", "Jimmy"),
        ("Architecture", "Jimmy"),
        ("Beta status + ask", "Aisling"),
        ("Close", "Both"),
    ]
    for section, who in rows:
        row = table.add_row().cells
        row[0].text = section
        row[1].text = who

    doc.add_heading("Demo backup (if they ask for live)", level=1)
    for step in [
        "Open Wisper → show welcome guide / model tier checkmarks",
        "Pick a short audio file → transcribe",
        'Show library search + "Transcription runs locally" copy',
    ]:
        doc.add_paragraph(step, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    shutil.copy2(OUT, OUT_DOWNLOADS)
    print(f"Wrote {OUT}")
    print(f"Copied to {OUT_DOWNLOADS}")


if __name__ == "__main__":
    main()
