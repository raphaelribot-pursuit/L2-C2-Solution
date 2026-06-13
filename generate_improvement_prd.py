"""Fill Week 2 Improvement PRD from Pursuit template, preserving layout/styles."""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BASE = Path(__file__).resolve().parent
SRC = Path(
    r"C:\Users\Aisling Ld Pursuit\Downloads\Aisling Improvment 20260515 PRD Template.docx"
)
WORK = BASE / "_improvement_prd_working.docx"
OUT = BASE / "Aisling Improvement Week 2 PRD - Filled.docx"

DIRECTION_FILL = "D6E4F0"

REPLACEMENTS = {
    "[BUILD NAME]": "WISPER — WEEK 2: GUIDED FIRST SCREEN",
    "[Build name]": "Wisper — Guided First Screen (Progressive Disclosure)",
    "[Your name(s)]": "Aisling Ld Pursuit",
    "[Date]": "June 8, 2026",
    "Product Requirements Document: Net New Build": (
        "Product Requirements Document: Week 2 Improvement"
    ),
    "[Write your problem description here.]": (
        "First-time Wisper users experience confusion and hesitation on the home screen "
        "during beta testing because Compute, Language, Record/import, URL import, and "
        "library controls are all visible at once before they complete a first transcription, "
        "resulting in slower time-to-first-transcript, abandoned sessions, and the impression "
        "that the app is harder to use than Whisper Notes despite strong backend capabilities."
    ),
    "[Insert data point or research finding]": (
        "Hick's Law: decision time increases with the number of choices — a core Laws of UX "
        "pattern from L2 Week 2 prioritization work."
    ),
    "[Insert user pain point observed or validated]": (
        "Beta testers must parse three panels (Compute, Language, Record or import) before "
        "understanding the single job: get audio in, transcribe locally."
    ),
    "[Insert market insight or competitive context]": (
        "Whisper Notes leads with a simple capture-first home screen; Wisper already matches "
        "on progress feedback during long jobs but not on first-screen clarity."
    ),
    "[Write your opportunity statement here.]": (
        "Reduce first-session friction so more beta testers complete one local transcription "
        "without reading the README — unlocking the value Wisper already delivers during "
        "download and transcribe progress."
    ),
    "[Insert market size or growth data]": (
        "Week 2 beta goal: trusted testers install, transcribe once, and return to the library — "
        "first-screen clarity directly affects activation for a small but high-signal cohort."
    ),
    "[Insert link to strategy doc or competitive analysis, if available]": (
        "See project CHANGELOG.md (Phase 1 exit QA), local ROADMAP.md, and Week 2 "
        "impact/effort prioritization (progressive disclosure vs release CI vs model download)."
    ),
    "[Who they are and what they care about]": (
        "Primary: beta testers and first-time users — students and professionals trying Wisper "
        "for the first time who want to transcribe one file or recording quickly without "
        "configuring GPU backends or optional URL import."
    ),
    "[Who else interacts with this product, if anyone]": (
        "Secondary: returning power users who want access to language, compute, and URL import "
        "without cluttering the default view."
    ),
    "[Write your solution description here.]": (
        "Wisper Week 2 improves the existing desktop app with a guided first screen and "
        "progressive disclosure. On launch, users see one clear job — drop or choose audio, "
        "or record — with Language defaulting to Auto-detect and Compute defaulting to the "
        "best available backend. Advanced options (language picker, CPU/GPU toggle, URL import, "
        "model path hints) move behind a collapsible Advanced section. An empty-state hero "
        "appears when no transcript is loaded, replacing the current three-panel wall. "
        "Long-job progress UI (Download → Transcribe, GPU fallback banners) stays unchanged."
    ),
    "[Write your value proposition here.]": (
        "Beta testers and privacy-conscious users who open Wisper for the first time use the "
        "guided first screen to start a local transcription in seconds. Unlike the current "
        "beta UI that exposes every workflow upfront, this improvement matches Whisper Notes' "
        "simplicity at the front door while keeping power features one click away."
    ),
    "[Write your vitamin value prop]": (
        "A clear, uncluttered home screen that tells users what to do first — the baseline "
        "expectation for any consumer-grade desktop app."
    ),
    "[Write your painkiller value prop]": (
        "Eliminates first-run overwhelm so users reach their first transcript instead of "
        "bouncing after scanning Compute, Language, and URL panels."
    ),
    "[Your steroid value prop]": (
        "Open Wisper, drop an MP3, tap Transcribe — the app feels obvious before the progress "
        "bar even moves, while advanced users can still expand GPU and URL tools when needed."
    ),
    "[What you’re trying to achieve — frame as an outcome, not a feature]": (
        "Increase first-session transcription completion among beta testers without removing "
        "existing power-user capabilities."
    ),
    "[Business or user benefit you’re optimizing for]": (
        "Time-to-first-successful-transcript and perceived ease of use at launch."
    ),
    "[Success outcome for this version]": (
        "A first-time user can complete record or file import → transcribe with no required "
        "configuration on the default view."
    ),
    "[What’s explicitly out of scope for this version, and why]": (
        "In-app model download, SRT export, release CI fixes, and security hardening — tracked "
        "on the broader wishlist; this PRD covers UI simplification only."
    ),
    "[What you’re NOT building right now]": (
        "Full settings page, redesign of library/transcript editor, diarization, or removing "
        "URL/GPU features — only hiding them by default."
    ),
    "[Goal from section 2b]": "First-session activation",
    "[Measurable indicator]": "User completes first transcription end-to-end",
    "[How you’ll track it]": "First-session completion rate (manual beta log + optional in-app event)",
    "[Specific number or threshold]": ">80% of first-time beta testers within first session",
    "[Add supporting links, research, wireframes, or open questions here.]": (
        "__APPENDIX_PLACEHOLDER__"
    ),
}

NEEDS = [
    "As a first-time beta tester, I need to see one obvious next step when I open Wisper because I should not have to read documentation to transcribe my first file.",
    "As a student importing a lecture, I need file pick and transcribe to be front and center because I do not care about GPU backends on day one.",
    "As a returning user, I need to expand Advanced options when I want URL import or a fixed language because I should not lose power features.",
    "As a privacy-conscious user, I need the simplified screen to still show that transcription is local because trust matters even in a minimal UI.",
]

GOALS = [
    "Reduce time from app open to first transcription attempt to under 30 seconds for file import.",
    "Keep all existing Phase 1 flows (record, file, URL, GPU toggle) accessible via Advanced — no feature removal.",
    "Preserve accessibility: progress regions, labels, and keyboard paths for primary actions.",
]

NON_GOALS = [
    "Replacing the first-run model setup banner with in-app download — separate improvement.",
    "Changing transcription engine, library schema, or export formats in this slice.",
]

METRICS = [
    (
        "First-session activation",
        "First-time user completes one transcription",
        "First-session completion rate",
        ">80% of beta testers",
    ),
    (
        "Time to first action",
        "User opens file picker or starts record",
        "Median seconds from launch to first action",
        "<30 seconds",
    ),
    (
        "Clarity (qualitative)",
        "Tester describes primary action without prompting",
        "Moderated beta feedback (5 users)",
        "4/5 can state next step in 10s",
    ),
    (
        "Power-user regression",
        "Advanced users still reach URL/GPU controls",
        "Task success in unmoderated test",
        "100% find Advanced within 2 clicks",
    ),
]

JOURNEYS = [
    {
        "title": "User Journey 1: First-time beta tester transcribes a local file",
        "context": (
            "Highest-frequency Week 2 path. Optimizing for clarity and speed to first "
            "transcript — the activation metric for beta."
        ),
        "subs": [
            (
                "Sub-journey: Landing on a clear home screen",
                [
                    ("P0", "User sees a guided empty state when no transcript is loaded (drop zone + primary actions)."),
                    ("P0", "User sees Record and Choose audio file as primary buttons without scrolling past Compute/Language panels."),
                    ("P0", "Advanced section is collapsed by default and labeled clearly (e.g. Advanced options)."),
                    ("P1", "User sees a one-line subtitle reinforcing local transcription (no cloud upload)."),
                    ("P2", "User can dismiss empty-state tips after first successful transcription."),
                ],
            ),
            (
                "Sub-journey: Completing first transcription with defaults",
                [
                    ("P0", "User can choose a file and tap Transcribe with Language set to Auto-detect and Compute on system default."),
                    ("P0", "User still sees two-step Download → Transcribe progress when applicable (unchanged)."),
                    ("P1", "User sees first-run model banner above or below hero if model is missing (existing onboarding preserved)."),
                    ("P2", "User can pin Advanced section open for future sessions."),
                ],
            ),
        ],
    },
    {
        "title": "User Journey 2: Returning user needs URL import or GPU control",
        "context": (
            "Progressive disclosure must not trap power users. Advanced paths stay "
            "discoverable without cluttering the default view."
        ),
        "subs": [
            (
                "Sub-journey: Expanding advanced options",
                [
                    ("P0", "User can expand Advanced to reveal Language select, Compute toggle, and URL import."),
                    ("P0", "Expanded state persists for the session (or until user collapses)."),
                    ("P1", "User sees hint text explaining when to change language or GPU (wrong language, GPU failure)."),
                    ("P2", "User can set a preference to always show Advanced (settings stub or localStorage)."),
                ],
            ),
            (
                "Sub-journey: URL import from Advanced",
                [
                    ("P0", "User can paste a URL and download → transcribe from the Advanced section."),
                    ("P0", "Download/transcribe progress and error phase tagging behave as in Phase 1."),
                    ("P1", "User sees yt-dlp setup hint in Advanced when binary is missing."),
                ],
            ),
        ],
    },
]

APPENDIX_LINES = [
    "Prioritization context (Week 2)",
    "• Impact/effort matrix + MoSCoW: Windows release CI is Must for installability; this UI improvement is Should — best impact/effort after CI.",
    "• Emotional pick aligned with framework #2: progressive disclosure, not CI (infra).",
    "",
    "Technical approach",
    "• Frontend-only change in wisper/src/App.tsx and styles — collapse Compute + Language + URL into Advanced; add empty-state hero.",
    "• No Rust/API changes required for MVP of this improvement.",
    "• Preserve aria-live progress blocks and first-run onboarding banner.",
    "",
    "Dependencies / blockers",
    "• Beta testers need an installable build (Windows release CI) to validate metrics — track separately.",
    "• Model must be present for transcription; in-app model download is a follow-up PRD.",
    "",
    "Open questions",
    "• Should Advanced remember expanded state across app restarts?",
    "• Mobile-style single FAB vs two primary buttons (Record / Choose file)?",
    "• Include a 3-step first-run coach mark overlay in Week 2 or defer?",
    "",
    "Success validation plan",
    "• 5 moderated beta sessions: can user transcribe one file without help?",
    "• Regression: power user finds URL import in <2 clicks.",
    "• Screenshot before/after for L2 submission.",
]


def paragraph_fill(p: Paragraph) -> str | None:
    pPr = p._element.pPr
    if pPr is None:
        return None
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        return None
    return (shd.get(qn("w:fill")) or "").upper()


def set_paragraph_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def replace_placeholders(p: Paragraph) -> None:
    text = p.text
    if not text:
        return
    new = text
    for old, val in REPLACEMENTS.items():
        if old in new:
            new = new.replace(old, val)
    if new != text:
        set_paragraph_text(p, new)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    el.getparent().remove(el)


def clone_after(ref: Paragraph, text: str, style=None) -> Paragraph:
    new_el = copy.deepcopy(ref._element)
    ref._element.addnext(new_el)
    new_p = Paragraph(new_el, ref._parent)
    set_paragraph_text(new_p, text)
    if style is not None:
        new_p.style = style
    return new_p


def fill_appendix(doc: Document) -> None:
    anchor = None
    for p in doc.paragraphs:
        if "__APPENDIX_PLACEHOLDER__" in (p.text or ""):
            anchor = p
            break
    if anchor is None:
        return
    set_paragraph_text(anchor, APPENDIX_LINES[0])
    insert = anchor
    for line in APPENDIX_LINES[1:]:
        insert = clone_after(insert, line, anchor.style)


def fill_user_needs(doc: Document) -> None:
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "Key User Needs":
            idx = i
            break
    if idx is None:
        return

    heading = doc.paragraphs[idx]
    to_delete: list[Paragraph] = []
    style_ref = None
    for p in doc.paragraphs[idx + 1 :]:
        if p.style and p.style.name.startswith("Heading"):
            break
        if (p.text or "").strip():
            if style_ref is None:
                style_ref = p.style
            to_delete.append(p)
    for p in reversed(to_delete):
        delete_paragraph(p)

    if style_ref is None:
        style_ref = heading.style

    insert = heading
    for need in NEEDS:
        insert = clone_after(insert, need, style_ref)


def fill_goals_and_non_goals(doc: Document) -> None:
    def fill_block(heading: str, lines: list[str]) -> None:
        idx = None
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() == heading:
                idx = i
                break
        if idx is None:
            return
        paras: list[Paragraph] = []
        for p in doc.paragraphs[idx + 1 :]:
            if p.style and p.style.name.startswith("Heading"):
                break
            if (p.text or "").strip():
                paras.append(p)
        if not paras:
            return
        set_paragraph_text(paras[0], lines[0])
        insert = paras[0]
        for i, line in enumerate(lines[1:], start=1):
            if i < len(paras):
                set_paragraph_text(paras[i], line)
            else:
                insert = clone_after(insert, line, paras[0].style)
        for extra in paras[len(lines) :]:
            delete_paragraph(extra)

    fill_block("Goals", GOALS)
    fill_block("Non-Goals", NON_GOALS)


def fill_metrics_table(doc: Document) -> None:
    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr != ["Goal", "Signal", "Metric", "Target"]:
            continue
        while len(table.rows) < len(METRICS) + 1:
            table.add_row()
        for ri, row in enumerate(METRICS, start=1):
            for ci, val in enumerate(row):
                cell = table.rows[ri].cells[ci]
                if cell.paragraphs:
                    set_paragraph_text(cell.paragraphs[0], val)
                else:
                    cell.text = val


def replace_requirements(doc: Document) -> None:
    req_idx = start_idx = end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "3. REQUIREMENTS":
            req_idx = i
        if start_idx is None and t.startswith("User Journey 1:"):
            start_idx = i
        if t.startswith("4. APPENDIX"):
            end_idx = i
            break
    if req_idx is None or end_idx is None:
        return

    if start_idx is not None and start_idx < end_idx:
        journey_style = doc.paragraphs[start_idx].style
        body_ref = doc.paragraphs[start_idx + 1]
        body_style = body_ref.style
        for i in range(end_idx - 1, start_idx - 1, -1):
            delete_paragraph(doc.paragraphs[i])
    else:
        journey_style = doc.paragraphs[req_idx].style
        body_style = doc.paragraphs[req_idx].style
        for i in range(end_idx - 1, req_idx, -1):
            delete_paragraph(doc.paragraphs[i])

    insert_after = doc.paragraphs[req_idx]
    for journey in JOURNEYS:
        insert_after = clone_after(insert_after, journey["title"], journey_style)
        insert_after = clone_after(
            insert_after, f"Context:  {journey['context']}", body_style
        )
        insert_after = clone_after(insert_after, "", body_style)
        for sub_title, reqs in journey["subs"]:
            insert_after = clone_after(insert_after, sub_title, body_style)
            for pri, req in reqs:
                insert_after = clone_after(
                    insert_after, f"[{pri}]  {req}", body_style
                )
            insert_after = clone_after(insert_after, "", body_style)


def remove_direction_boxes(doc: Document) -> None:
    to_delete = [p for p in doc.paragraphs if paragraph_fill(p) == DIRECTION_FILL]
    for p in reversed(to_delete):
        delete_paragraph(p)


def remove_how_to_use_block(doc: Document) -> None:
    """Remove template instructions block after title."""
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "HOW TO USE THIS TEMPLATE":
            start = i
        if start is not None and t.startswith("Build name:"):
            end = i
            break
    if start is None or end is None:
        return
    for i in range(end - 1, start - 1, -1):
        delete_paragraph(doc.paragraphs[i])


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Template not found: {SRC}")
    try:
        shutil.copy2(SRC, WORK)
    except PermissionError:
        if not WORK.exists():
            raise

    doc = Document(str(WORK))

    for p in doc.paragraphs:
        replace_placeholders(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders(p)

    fill_user_needs(doc)
    fill_goals_and_non_goals(doc)
    fill_metrics_table(doc)
    fill_appendix(doc)
    remove_how_to_use_block(doc)
    remove_direction_boxes(doc)
    replace_requirements(doc)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
